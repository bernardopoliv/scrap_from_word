import os

import docx
import openpyxl

fields_and_values = {}
output_filepath = 'C:\\Users\BERNARDO\Anaconda3\envs\SCRAP_FROM_WORD\\Output.xlsx'
replace_filepath = 'C:\\Users\BERNARDO\Anaconda3\envs\SCRAP_FROM_WORD\\replaced_output.xlsx'


def lookup(doc, fields_and_values, paragraph_index: int, runs_index: int, variable_name: str):
    value = doc.paragraphs[paragraph_index].runs[runs_index].text
    fields_and_values.update({variable_name: value})

    return value


def get_allparagraph(doc, fields_and_values, paragraph_index: int, variable_name: str):
    value = doc.paragraphs[paragraph_index].text
    fields_and_values.update({variable_name: value})
    return value


def look_value(doc, paragraph_index: int, runs_index: int):
    value = doc.paragraphs[paragraph_index].runs[runs_index].text
    return value


def export(fields_and_values, output_file=None):
    output_wb = openpyxl.Workbook()
    output_sheet = output_wb.active

    row = 1
    for key, value in fields_and_values.items():
        output_sheet.cell(row=row, column=1).value = key
        output_sheet.cell(row=row, column=2).value = value
        row += 1

    if output_file:
        output_wb.save(output_file)
    else:
        output_wb.save(output_filepath)


def scan():
    i = 0
    for p in doc.paragraphs:
        print(i, p.text)
        i += 1


def find_value(paragraph_index: int):
    for p in range(0, 100):
        try:
            print(str(p) + doc.paragraphs[paragraph_index].runs[p].text)
        except:
            pass


def replace_values(fields_and_values, replace_filename: str, doc):
    replace_wb = openpyxl.load_workbook(replace_filename)
    replace_sheet = replace_wb['Sheet']

    filled_field = 0

    for row in range(1,200):
        if replace_sheet.cell(row=row, column=2).value != None:
            filled_field += 1

    for row in range(1,filled_field):
        new_value = str(replace_sheet.cell(row=row, column=3).value)
        replace_sheet.cell(row=row, column=2).value = new_value
        replace_sheet.cell(row=row, column=3).value = None

    for row in range(1,filled_field):
        fields_and_values.update({replace_sheet.cell(row=row, column=1).value: replace_sheet.cell(row=row, column=2).value})

    return fields_and_values

def generate_new_doc(document_type: str, doc, fields_and_values):

    if document_type == 'CO':

        # COMMERCIAL TEMPLATE
        # PAGE 1
        doc.paragraphs[0].runs[0].text = fields_and_values['PERMIT_TYPE']
        doc.paragraphs[2].text = fields_and_values['AGREEMENT_DATE']
        doc.paragraphs[6].runs[3].text = fields_and_values['LANDLORD']
        doc.paragraphs[16].runs[3].text = fields_and_values['TENANT_NAME']
        doc.paragraphs[28].runs[3].text = fields_and_values['LOT_NUMBER']
        doc.paragraphs[28].runs[8].text = fields_and_values['PLANNO']


        # PAGE 3
        doc.paragraphs[57].runs[1].text = fields_and_values['TENANT_LEASE_TERM']
        doc.paragraphs[57].runs[4].text = fields_and_values['INITIAL_DATE']
        doc.paragraphs[57].runs[7].text = fields_and_values['EXPIRY_DATE']
        doc.paragraphs[59].runs[1].text = fields_and_values['TENANT_PREMISE_PURPOSE']

        # PAGE 4
        doc.paragraphs[65].runs[2].text = fields_and_values['RENT_EXPIRE_DATE']
        doc.paragraphs[65].runs[6].text = fields_and_values['YEARLY_RENT_AMOUNT']

        doc.paragraphs[65].runs[9].text = fields_and_values['YEARLY_RENT_PAY']

        doc.paragraphs[65].runs[12].text = fields_and_values['MONTH_4_PAYMENT_AGREEMENT']
        doc.paragraphs[65].runs[15].text = fields_and_values['TENANT_FISCAL_START_DATE']
        doc.paragraphs[65].runs[18].text = fields_and_values['TENANT_FISCAL_END_DATE']
        doc.paragraphs[65].runs[21].text = fields_and_values['TENANT_CONFIRMED_PAYMENT_DATE']


        # PAGE 7
        # fields_and_values['TENANT_CONIFIRM_PAY_WATER'] = 'YES'
        # fields_and_values['TENANT_CONIFIRM_PAY_GAS'] = 'YES'
        # fields_and_values['TENANT_CONIFIRM_PAY_TELEPHONE'] = 'YES'
        # fields_and_values['TENANT_CONIFIRM_PAY_LIGHT'] = 'YES'
        # fields_and_values['TENANT_CONIFIRM_PAY_POWER'] = 'YES'
        # fields_and_values['TENANT_CONIFIRM_PAY_HEAT'] = 'YES'
        # fields_and_values['TENANT_CONIFIRM_PAY_AIRCONDITIONING'] = 'YES'
        # fields_and_values['TENANT_CONIFIRM_PAY_SEWER'] = 'YES'
        # fields_and_values['TENANT_CONIFIRM_PAY_GARBAGE'] = 'YES'


        # PAGE 10
        doc.paragraphs[134].runs[1].text = fields_and_values['TENANT_LIABILITY_INSURANCE_AMOUNT']

        # PAGE 19
        doc.paragraphs[233].runs[2].text = fields_and_values['TENANT_NAME_NOTICE']

        # PAGE 20
        doc.paragraphs[253].runs[5].text = fields_and_values['CHIEF_LAND_COMM_LEASE']
        doc.paragraphs[255].runs[5].text = fields_and_values['COUNCILLOR1_LAND_COMM_LEASE']
        doc.paragraphs[257].runs[5].text = fields_and_values['COUNCILLOR2_LAND_COMM_LEASE']
        # doc.paragraphs[260].runs[6].text = fields_and_values['WITNESS1_LAND_COMM_LEASE']


        doc.save('test.docx')

def get_values(document_type: str, doc: docx.Document):  # 'AG' or 'CO' or 'RE'
    fields_and_values = {}
    if document_type == 'CO':

        # COMMERCIAL TEMPLATE
        # PAGE 1
        lookup(doc, fields_and_values, 0, 0, 'PERMIT_TYPE')
        get_allparagraph(doc, fields_and_values, 2, 'AGREEMENT_DATE')
        lookup(doc, fields_and_values, 6, 3, 'LANDLORD')
        lookup(doc, fields_and_values, 16, 3, 'TENANT_NAME')
        lookup(doc, fields_and_values, 28, 3, 'LOT_NUMBER')
        lookup(doc, fields_and_values, 28, 8, 'PLANNO')

        # PAGE 3
        lookup(doc, fields_and_values, 57, 1, 'TENANT_LEASE_TERM')
        lookup(doc, fields_and_values, 57, 4, 'INITIAL_DATE')  # TODO: check if it is missing in original variables
        lookup(doc, fields_and_values, 57, 7, 'EXPIRY_DATE')
        lookup(doc, fields_and_values, 59, 1, 'TENANT_PREMISE_PURPOSE')

        # PAGE 4
        lookup(doc, fields_and_values, 65, 2, 'RENT_EXPIRE_DATE')
        lookup(doc, fields_and_values, 65, 6, 'YEARLY_RENT_AMOUNT')

        if str('year') is str(doc.paragraphs[65].runs[9].text):
            fields_and_values.update({'YEARLY_RENT_PAY': 'YES'})  # TODO: ERROR
        else:
            fields_and_values.update({'YEARLY_RENT_PAY': 'NO'})

        lookup(doc, fields_and_values, 65, 12, 'MONTH_4_PAYMENT_AGREEMENT')
        lookup(doc, fields_and_values, 65, 15, 'TENANT_FISCAL_START_DATE')
        lookup(doc, fields_and_values, 65, 18, 'TENANT_FISCAL_END_DATE')
        lookup(doc, fields_and_values, 65, 21, 'TENANT_CONFIRMED_PAYMENT_DATE')

        # PAGE 7
        fields_and_values['TENANT_CONIFIRM_PAY_WATER'] = 'YES'
        fields_and_values['TENANT_CONIFIRM_PAY_GAS'] = 'YES'
        fields_and_values['TENANT_CONIFIRM_PAY_TELEPHONE'] = 'YES'
        fields_and_values['TENANT_CONIFIRM_PAY_LIGHT'] = 'YES'
        fields_and_values['TENANT_CONIFIRM_PAY_POWER'] = 'YES'
        fields_and_values['TENANT_CONIFIRM_PAY_HEAT'] = 'YES'
        fields_and_values['TENANT_CONIFIRM_PAY_AIRCONDITIONING'] = 'YES'
        fields_and_values['TENANT_CONIFIRM_PAY_SEWER'] = 'YES'
        fields_and_values['TENANT_CONIFIRM_PAY_GARBAGE'] = 'YES'

        # PAGE 10
        lookup(doc, fields_and_values, 134, 1, 'TENANT_LIABILITY_INSURANCE_AMOUNT')

        # PAGE 19
        lookup(doc, fields_and_values, 233, 2, 'TENANT_NAME_NOTICE')

        # PAGE 20
        lookup(doc, fields_and_values, 253, 5, 'CHIEF_LAND_COMM_LEASE')
        lookup(doc, fields_and_values, 255, 5, 'COUNCILLOR1_LAND_COMM_LEASE')
        lookup(doc, fields_and_values, 257, 5, 'COUNCILLOR2_LAND_COMM_LEASE')
        lookup(doc, fields_and_values, 260, 6, 'WITNESS1_LAND_COMM_LEASE')

        print(len(fields_and_values.keys()))
        export(fields_and_values)
        return fields_and_values

    elif document_type == 'RE':

        # RESIDENTIAL TEMPLATE
        fields_and_values = {}

        # PAGE 1
        lookup(doc, fields_and_values, 0, 0, 'PERMIT_TYPE')
        lookup(doc, fields_and_values, 1, 1, 'RESIDENTIAL_AGREEMENT_DATE')
        lookup(doc, fields_and_values, 15, 3, 'RESIDENTIAL_TENANT_NAME')

        # PAGE 4
        lookup(doc, fields_and_values, 62, 4, 'RESIDENTIAL_LEASE_INITIAL_DATE')  # TODO: put year together (runs[5)
        lookup(doc, fields_and_values, 62, 7, 'RESIDENTIAL_LEASE_EXPIRY_DATE')  # TODO: put year together (runs[8)

        # PAGE 7
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONFIRM_PAY_TAXES'] = 'YES'
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONFIRM_PAY_WATER'] = 'YES'
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONFIRM_PAY_GAS'] = 'YES'
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONFIRM_PAY_TELEPHONE'] = 'YES'
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONFIRM_PAY_LIGHT'] = 'YES'
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONFIRM_PAY_POWER'] = 'YES'
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONFIRM_PAY_HEAT'] = 'YES'
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONFIRM_PAY_AIRCONDITIONING'] = 'YES'
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONFIRM_PAY_SEWER'] = 'YES'
        fields_and_values['RESIDENTIAL_LEASE_TENANT_CONIFIRM_PAY_GARBAGE'] = 'YES'

        # PAGE 18 TODO: check better this part with client
        lookup(doc, fields_and_values, 219, 2, 'RESIDENTIAL_LEASE_TENANT_NAME_NOTICE')
        lookup(doc, fields_and_values, 1, 1, 'RESIDENTIAL_LEASE_TENANT_NAME-NOTICE_DELIVERY: MAIL / EMAIL')
        lookup(doc, fields_and_values, 15, 3, 'RESIDENTIAL_LEASE_TENANT_NAME_MAILING_ADDRESS')
        lookup(doc, fields_and_values, 15, 3, 'RESIDENTIAL_LEASE_TENANT_NAME_EMAIL')

        # PAGE 20 TODO: check better this part with client
        lookup(doc, fields_and_values, 216, 2, 'CHIEF_RESIDENTIAL_LEASE')
        lookup(doc, fields_and_values, 1, 1, 'COUNCILLOR1_RESIDENTIAL_LEASE')
        lookup(doc, fields_and_values, 15, 3, 'COUNCILLOR2_RESIDENTIAL_LEASE')
        lookup(doc, fields_and_values, 15, 3, 'WITNESS1_RESIDENTIAL_LEASE')

        # PAGE 20 TODO: check better this part with client
        lookup(doc, fields_and_values, 216, 2, 'TENANT_NAME')
        lookup(doc, fields_and_values, 1, 1, 'AFFIDAVIT_WITNESS_RESIDENTIAL_LEASE_DATE')
        lookup(doc, fields_and_values, 15, 3, 'AFFADAVIT_WITNESS_RESIDENTIAL_LEASE_SIGNED')
        lookup(doc, fields_and_values, 15, 3, 'AFFADAVIT_WITNESS_RESIDENTIAL_LEASE_NOTARY_SIGNED')

        print(len(fields_and_values.keys()))
        export(fields_and_values)

        return fields_and_values

    elif document_type == 'AG':

        # AGRICUTURAL TEMPLATE
        fields_and_values = {}

        AG_PERMIT_TYPE = doc.paragraphs[1].text
        AG_PERMIT_TYPE = AG_PERMIT_TYPE[1:]
        fields_and_values.update({'AG_PERMIT_TYPE': AG_PERMIT_TYPE})

        AG_PERM_PRELIM_START_DATE = str(
            lookup(doc, fields_and_values, 4, 3, 'AG_PERM_PRELIM_START_DATE') + lookup(doc, fields_and_values, 4, 5,
                                                                                       'AG_PERM_PRELIM_START_DATE'))
        AG_PERM_PRELIM_START_DATE = AG_PERM_PRELIM_START_DATE.replace(' day of', ',')
        fields_and_values['AG_PERM_PRELIM_START_DATE'] = AG_PERM_PRELIM_START_DATE

        lookup(doc, fields_and_values, 7, 2, "AG_PERM_GRANTOR")
        lookup(doc, fields_and_values, 9, 3, "AG_PERMIT_GRANTEE")
        lookup(doc, fields_and_values, 20, 1, "AG_LOCATION")
        lookup(doc, fields_and_values, 20, 5, "AG_AREA_HECATARES")
        AG_AREA_ACRES = str(lookup(doc, fields_and_values, 20, 8, "AG_AREA_ACRES"))
        fields_and_values['AG_AREA_ACRES'] = AG_AREA_ACRES[:-1]

        lookup(doc, fields_and_values, 20, 11, "AG_LAND_USE")

        # PAGE 2
        AG_PERMIT_START_DATE = str(
            lookup(doc, fields_and_values, 23, 8, 'AG_PERMIT_START_DATE') + lookup(doc, fields_and_values, 23, 10,
                                                                                   'AG_PERMIT_START_DATE'))
        fields_and_values['AG_PERMIT_START_DATE'] = AG_PERMIT_START_DATE

        AG_PERMIT_END_DATE = str(
            lookup(doc, fields_and_values, 23, 15, 'AG_PERMIT_END_DATE') + lookup(doc, fields_and_values, 23, 17,
                                                                                  'AG_PERMIT_END_DATE'))
        fields_and_values['AG_PERMIT_END_DATE'] = AG_PERMIT_END_DATE

        AG_PERMIT_COMMENT = str(
            lookup(doc, fields_and_values, 31, 4, 'AG_PERMIT_COMMENT') + lookup(doc, fields_and_values, 31, 5,
                                                                                'AG_PERMIT_COMMENT') + lookup(doc,
                                                                                                              fields_and_values,
                                                                                                              31, 6,
                                                                                                              'AG_PERMIT_COMMENT'))
        fields_and_values.update({'AG_PERMIT_COMMENT': AG_PERMIT_COMMENT})

        # PAGE 3
        lookup(doc, fields_and_values, 0, 0, 'AG_LAND_CHEMICAL_PRESENT')
        lookup(doc, fields_and_values, 0, 0, 'AG_LAND_CHEMICAL_TYPE')

        AG_ENVIRONMENTAL_ISSUE = doc.paragraphs[64].text
        AG_ENVIRONMENTAL_ISSUE = AG_ENVIRONMENTAL_ISSUE[4:]
        fields_and_values.update({'AG_ENVIRONMENTAL_ISSUE': AG_ENVIRONMENTAL_ISSUE})  # TODO: QUESTION

        lookup(doc, fields_and_values, 0, 0, 'AG_ENVIRONMENTAL_ISSUE_TYPE')
        lookup(doc, fields_and_values, 0, 0, 'AG_LAND_CHEMICAL_PRESENT')

        AG_CULTURAL_SIG_FOUND = get_allparagraph(doc, fields_and_values, 68, 'AG_CULTURAL_SIG_FOUND')
        AG_CULTURAL_SIG_FOUND = AG_CULTURAL_SIG_FOUND[4:]
        fields_and_values.update({'AG_CULTURAL_SIG_FOUND': AG_CULTURAL_SIG_FOUND})

        AG_CULTURAL_SIG_FOUND_TYPE = get_allparagraph(doc, fields_and_values, 68, 'AG_CULTURAL_SIG_FOUND_TYPE')
        AG_CULTURAL_SIG_FOUND_TYPE = AG_CULTURAL_SIG_FOUND_TYPE[4:]
        fields_and_values.update({'AG_CULTURAL_SIG_FOUND_TYPE': AG_CULTURAL_SIG_FOUND_TYPE})

        # PAGE 4
        AG_ADDRESS_NAME = look_value(doc, 86, 6)
        AG_ADDRESS_NAME = AG_ADDRESS_NAME[:-1]
        fields_and_values.update({'AG_ADDRESS_NAME': AG_ADDRESS_NAME})
        lookup(doc, fields_and_values, 87, 7, 'AG_ADDRESS_STREET')
        lookup(doc, fields_and_values, 88, 7, 'AG_ADDRESS_CITY')
        lookup(doc, fields_and_values, 88, 10, 'AG_ADDRESS_PROVINCE')
        lookup(doc, fields_and_values, 89, 7, 'AG_ADDRESS_POSTAL_CODE')

        # PAGE 5 (NO DATA)
        # PAGE 6
        lookup(doc, fields_and_values, 170, 0, 'AG_SIGNED_CHIEF')
        lookup(doc, fields_and_values, 171, 0, 'AG_SIGNED_CHIEF_NAME')
        lookup(doc, fields_and_values, 175, 0, 'AG_SIGNED_COUNCILLOR1')
        lookup(doc, fields_and_values, 176, 0, 'AG_SIGNED_COUNCILLOR1_NAME')
        lookup(doc, fields_and_values, 180, 0, 'AG_SIGNED_COUNCILLOR2')
        lookup(doc, fields_and_values, 181, 0, 'AG_SIGNED_COUNCILLOR2_NAME')
        lookup(doc, fields_and_values, 141, 0, 'AG_SIGNED_WITNESS')
        lookup(doc, fields_and_values, 141, 4, 'AG_SIGNED_PERMITTEE_NAME')
        lookup(doc, fields_and_values, 181, 0, 'AG_SIGNED_COUNCILLOR2_NAME')
        lookup(doc, fields_and_values, 181, 0, 'AG_SIGNED_COUNCILLOR2_NAME')

        print(len(fields_and_values.keys()))

        export(fields_and_values)

        return fields_and_values




